#############
# IMPORTS
#############

import re  # Library for regular expressions and pattern matching
import docx  # Library for parsing Word documents
from docx import Document  # Specifically import Document class from docx

#############
# MODEL LOADING
#############

import spacy  # Library for advanced text analysis and NLP
from spacy.cli import download

# Global variable to hold the loaded spaCy model
nlp = None

def get_spacy_model(model_name="en_core_web_sm"):
    """Lazily loads or downloads the specified spaCy model."""
    global nlp
    if nlp is None:
        try:
            nlp = spacy.load(model_name)
        except IOError:
            print(f"Model {model_name} not found, downloading...")
            download(model_name)
            nlp = spacy.load(model_name)
    return nlp

#############
# KEYWORDS DICTIONARY
#############

SECTION_KEYWORDS = {  # Dictionary to map common section names to keywords
    "HEADER": ["name", "contact"],
    "SUMMARY": ["summary", "objective"],
    "SKILLS": ["skills", "expertise", "abilities"],
    "EXPERIENCE": ["experience", "work", "positions"],
    "EDUCATION": ["education", "qualifications", "degrees"],
}

#############
# FUNCTION: IDENTIFY SECTIONS
#############

def identify_sections(file_path):
    """
    Identifies sections in a Word document based on keywords and structure.
    """
    sections = {}
    with Document(file_path) as doc:
        for paragraph in doc.paragraphs:
            doc = nlp(paragraph.text)
            for section_name, keywords in SECTION_KEYWORDS.items():
                if any(token.text.lower() in keywords for token in doc):
                    sections[section_name] = sections.get(section_name, "") + paragraph.text + "\n"
    return sections


def suggest_resume_formatting(text):
    """
    Suggests formatting improvements for a resume based on text analysis.
    """
    suggestions = []

    # Rule for bullet points
    if text.count('-') > 5 and text.count('*') == 0:
        suggestions.append("Consider using bullet points ('*') for listing achievements or responsibilities for consistency.")

    # Rule for headings
    if "Experience" not in text or "Education" not in text:
        suggestions.append("Ensure key sections like 'Experience' and 'Education' are clearly marked with headings.")

    # Rule for contact information placement
    if "Contact" not in text[:100]:  # Assuming contact info should be at the top
        suggestions.append("Place contact information prominently at the top of your resume.")
    
    # Rule for consistent date format
    date_pattern = r'\b\d{1,2}/\d{4}\b|\b[a-zA-Z]+\s\d{4}\b'
    if not re.search(date_pattern, text):
        suggestions.append("Ensure a consistent date format (e.g., 'MM/YYYY' or 'Month Year') throughout the resume.")
    
    # Rule for action verbs
    action_verbs = ["managed", "achieved", "developed", "implemented", "led", "analyzed", "created"]
    used_action_verbs = [verb for verb in action_verbs if verb in text.lower()]
    
    if len(used_action_verbs) == 0:
        suggestions.append("Consider using action verbs like 'managed,' 'achieved,' 'developed,' etc., to describe your experiences.")
    
    # Rule for consistent font size and style (placeholder)
    # You may need to use external libraries or tools to check font consistency.
    
    # Rule for white space (placeholder)
    # Implement logic to check for consistent white space and spacing between sections.
    
    # Rule for quantifiable achievements (placeholder)
    # Implement logic to check for quantifiable achievements.
    
    # Rule for keywords (placeholder)
    # Implement logic to check for relevant keywords.
    
    # Rule for consistent formatting of section titles
    section_titles = ["Professional Summary", "Skills", "Experience", "Education", "Certifications"]
    section_pattern = r'\b(?:' + '|'.join(section_titles) + r')\b'
    mismatched_titles = [title for title in section_titles if re.search(r'\b' + title + r'\b', text, flags=re.IGNORECASE) is None]
    
    if mismatched_titles:
        suggestions.append(f"Ensure consistent formatting of section titles like: {', '.join(mismatched_titles)}")
    
    # Rule for spelling and grammar (placeholder)
    # Implement logic to check for spelling and grammar errors.
    
    # Rule for consistent bullet point symbols (placeholder)
    # Implement logic to check for consistent bullet point symbols.
    
    # Rule for length (placeholder)
    # Implement logic to check for resume length and suggest trimming if necessary.
    
    return suggestions



def analyze_text_hierarchy(file_path):
    """
    Analyzes the hierarchy of text elements in a resume.
    """
    text_hierarchy = {"bold": 0, "headings": 0, "italics": 0}
    with Document(file_path) as doc:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.bold: text_hierarchy["bold"] += 1
                if run.font.underline: text_hierarchy["italics"] += 1
                if run.style and run.style.name.startswith("Heading"): text_hierarchy["headings"] += 1
    return text_hierarchy

def check_visual_elements(file_path):
    """
    Checks for consistency and readability in font sizes, colors, and spacing.
    """
    visual_elements = {"font_size_consistency": True, "color_consistency": True}
    # Analyze font sizes and colors (implementation depends on document format)
    return visual_elements

def suggest_structure_improvements(sections):
    """
    Suggests improvements based on the structure of identified sections.
    """
    suggestions = []
    # Example: check for consistent use of headings and subheadings
    if "EXPERIENCE" in sections and sections["EXPERIENCE"].count('\n') > 5:
        suggestions.append("Consider using subheadings in the 'Experience' section for better readability.")
    # Add more suggestions based on section analysis
    return suggestions

def suggest_visual_improvements(font_sizes, colors, text_hierarchy):
    """
    Suggests improvements for visual elements like font sizes, colors, and spacing.
    """
    suggestions = []
    # Example: suggest larger font size if too small
    if min(font_sizes) < 10:
        suggestions.append("Increase font size for better readability.")
    # Add more suggestions based on visual element analysis
    return suggestions

# Example usage of the functions
# file_path = "path/to/resume.docx"
# sections = identify_sections(file_path)
# format_suggestions = suggest_resume_formatting("Your resume text here")
# print(format_suggestions)
